from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from werkzeug.exceptions import RequestEntityTooLarge
import os
import json
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import cohere
import logging
import bleach
from werkzeug.utils import secure_filename

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s %(levelname)s %(name)s %(message)s'
)
logger = logging.getLogger(__name__)

load_dotenv()

# Load environment-specific .env file
env = os.getenv("ENVIRONMENT", "local")
if env == "production":
    load_dotenv(".env.production")
else:
    load_dotenv(".env")

# Configure Cohere
COHERE_API_KEY = os.getenv("COHERE_API_KEY") 
if not COHERE_API_KEY:
    raise ValueError("COHERE_API_KEY environment variable is required")

co = cohere.Client(COHERE_API_KEY)
MODEL_NAME = "command-a-03-2025"
FALLBACK_MODEL = "command-r" 

app = Flask(__name__)

# Security Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'dev-key-change-in-production')

# CORS Configuration - restrict in production
if env == "production":
    allowed_origins = os.getenv('ALLOWED_ORIGINS', '').split(',')
    CORS(app, origins=allowed_origins)
else:
    CORS(app)  # Allow all origins in development

# Rate Limiter
limiter = Limiter(
    app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def sanitize_text(text):
    """Sanitize text input to prevent XSS attacks"""
    if not text:
        return ""
    # Remove HTML tags and limit length
    sanitized = bleach.clean(text, tags=[], strip=True)
    return sanitized[:50000]  # Limit to 50k characters

def validate_personal_info(personal_info):
    """Validate and sanitize personal information"""
    if not isinstance(personal_info, dict):
        return {}
    
    allowed_fields = [
        'fullName', 'address', 'phone', 'email', 'hiringManagerName',
        'hiringManagerTitle', 'companyName', 'companyAddress', 
        'positionTitle', 'howHeardAbout'
    ]
    
    sanitized = {}
    for field in allowed_fields:
        if field in personal_info:
            value = str(personal_info[field]).strip()
            if len(value) > 500:  # Limit field length
                value = value[:500]
            sanitized[field] = bleach.clean(value, tags=[], strip=True)
    
    return sanitized

@app.errorhandler(413)
def too_large(e):
    return jsonify({"error": "File too large. Maximum size is 16MB."}), 413

@app.errorhandler(RequestEntityTooLarge)
def handle_file_too_large(e):
    return jsonify({"error": "File too large. Maximum size is 16MB."}), 413

@app.errorhandler(429)
def ratelimit_handler(e):
    return jsonify({"error": "Rate limit exceeded. Please try again later."}), 429

@app.errorhandler(500)
def internal_error(e):
    logger.error(f"Internal server error: {str(e)}")
    return jsonify({"error": "Internal server error"}), 500

@app.route("/improve_resume", methods=["POST"])
@limiter.limit("10 per hour")  # Limit AI requests
def improve_resume():
    try:
        file = request.files.get("file")
        job_desc = request.form.get("jobDescription", "")

        if not file:
            return jsonify({"error": "No file uploaded"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"error": "Invalid file type. Only .doc and .docx files are allowed."}), 400

        if not job_desc.strip():
            return jsonify({"error": "Job description is required"}), 400

        # Sanitize inputs
        job_desc = sanitize_text(job_desc)
        
        if len(job_desc) < 50:
            return jsonify({"error": "Job description too short. Please provide a detailed job description."}), 400

        # Read Word document with error handling
        try:
            doc = Document(file)
            resume_text = "\n".join([p.text for p in doc.paragraphs])
            
            if len(resume_text.strip()) < 100:
                return jsonify({"error": "Resume content appears to be too short or empty."}), 400
                
        except Exception as e:
            logger.error(f"Error reading document: {str(e)}")
            return jsonify({"error": "Unable to read the uploaded document. Please ensure it's a valid .doc or .docx file."}), 400

        # Limit resume text length
        resume_text = resume_text[:20000]  # Limit to 20k characters

        # Prompt
        prompt = f"""Improve and rewrite this resume for the following job:

Job Description: 
{job_desc}

Current Resume:
{resume_text}

Please provide an improved version that:
1. Better aligns with the job requirements
2. Uses strong action verbs
3. Quantifies achievements where possible
4. Maintains professional formatting
5. Highlights relevant skills and experience

Return only the improved resume content."""

        try:
            response = co.chat(
                model=MODEL_NAME,
                message=prompt,
                max_tokens=4000  # Limit response length
            )
            improved_text = response.text
        except Exception as e:
            logger.error(f"Error generating content with primary model: {e}")
            try:
                # Fallback to secondary model
                response = co.chat(
                    model=FALLBACK_MODEL,
                    message=prompt,
                    max_tokens=4000
                )
                improved_text = response.text
            except Exception as fallback_e:
                logger.error(f"Error with fallback model: {fallback_e}")
                return jsonify({"error": "Failed to generate improved resume"}), 500

        # Create new Word document
        try:
            out_doc = Document()
            for line in improved_text.split("\n"):
                if line.strip():
                    out_doc.add_paragraph(line.strip())
            
            byte_io = BytesIO()
            out_doc.save(byte_io)
            byte_io.seek(0)

            return send_file(
                byte_io,
                download_name="Improved_Resume.docx",
                as_attachment=True,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            logger.error(f"Error creating document: {str(e)}")
            return jsonify({"error": "Failed to create document"}), 500
    
    except Exception as e:
        logger.error(f"Error in improve_resume: {e}")
        return jsonify({"error": "Internal server error"}), 500


@app.route("/generate_cover_letter", methods=["POST"])
@limiter.limit("15 per hour")  # Limit AI requests
def generate_cover_letter():
    try:
        file = request.files.get("file")
        job_desc = request.form.get("jobDescription", "")
        personal_info_str = request.form.get("personalInfo", "{}")
        
        if not file:
            return jsonify({"error": "No resume file uploaded"}), 400
        
        if not allowed_file(file.filename):
            return jsonify({"error": "Invalid file type. Only .doc and .docx files are allowed."}), 400
            
        if not job_desc.strip():
            return jsonify({"error": "Job description is required"}), 400
        
        # Sanitize inputs
        job_desc = sanitize_text(job_desc)
        
        if len(job_desc) < 50:
            return jsonify({"error": "Job description too short. Please provide a detailed job description."}), 400
        
        try:
            personal_info_raw = json.loads(personal_info_str)
            personal_info = validate_personal_info(personal_info_raw)
        except json.JSONDecodeError:
            logger.warning("Invalid JSON in personalInfo, using empty dict")
            personal_info = {}

        # Read resume with error handling
        try:
            doc = Document(file)
            resume_text = "\n".join([p.text for p in doc.paragraphs])
            
            if len(resume_text.strip()) < 100:
                return jsonify({"error": "Resume content appears to be too short or empty."}), 400
                
        except Exception as e:
            logger.error(f"Error reading document: {str(e)}")
            return jsonify({"error": "Unable to read the uploaded document. Please ensure it's a valid .doc or .docx file."}), 400

        # Limit resume text length
        resume_text = resume_text[:20000]  # Limit to 20k characters

        position_title = personal_info.get('positionTitle', '').strip() or "the position"

        prompt = f"""Create a professional cover letter for the {position_title} position. Based on the following information:
        
Job Description: 
{job_desc}

Resume Content: 
{resume_text}

Instructions:
1. Write ONLY the body content of the cover letter (no header, no contact information, no date)
2. Start with "Dear [Hiring Manager Name]," or "Dear Hiring Manager," 
3. Write 2-3 paragraphs:
   - Opening: Express interest in the specific position 
   - Body: Highlight relevant skills and experiences from the resume that match job requirements with specific examples
   - Closing: Express enthusiasm, mention next steps, and thank them
4. End with "Sincerely," followed by a blank line for signature
5. Use professional but engaging language
6. Do NOT include any placeholder brackets like [Your Name] or contact information
7. Do NOT include any header information, addresses, or dates
8. Focus on connecting resume experience to job requirements
9. Keep it concise and impactful

Additional context:
- Company: {personal_info.get('companyName', 'your company')}
- How heard about position: {personal_info.get('howHeardAbout', 'through your website')}"""

        try:
            response = co.chat(
                model=MODEL_NAME,
                message=prompt,
                max_tokens=3000  # Limit response length
            )
            cover_letter_content = response.text.strip()
        except Exception as e:
            logger.error(f"Error generating cover letter with primary model: {e}")
            try:
                # Fallback to secondary model
                response = co.chat(
                    model=FALLBACK_MODEL,
                    message=prompt,
                    max_tokens=3000
                )
                cover_letter_content = response.text.strip()
            except Exception as fallback_e:
                logger.error(f"Error with fallback model: {fallback_e}")
                return jsonify({"error": "Failed to generate cover letter"}), 500

        # Build Word document
        try:
            doc = Document()
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            if personal_info.get('fullName'):
                name_para = doc.add_paragraph(personal_info['fullName'])
                name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in name_para.runs:
                    run.bold = True
            
            contact_info = []
            if personal_info.get('address'):
                contact_info.append(personal_info['address'])
            if personal_info.get('phone') and personal_info.get('email'):
                contact_info.append(f"{personal_info['phone']} | {personal_info['email']}")
            elif personal_info.get('phone'):
                contact_info.append(personal_info['phone'])
            elif personal_info.get('email'):
                contact_info.append(personal_info['email'])
            
            for info in contact_info:
                p = doc.add_paragraph(info)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            doc.add_paragraph()
            date_p = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
            date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            doc.add_paragraph()

            recipient_added = False
            for field in ["hiringManagerName", "hiringManagerTitle", "companyName", "companyAddress"]:
                if personal_info.get(field):
                    doc.add_paragraph(personal_info[field])
                    recipient_added = True
            if recipient_added:
                doc.add_paragraph()

            # Process AI-generated content
            lines = cover_letter_content.split("\n")
            processed_lines = []
            for line in lines:
                line = line.strip()
                if line:
                    if line.startswith('Dear ') and '[Hiring Manager Name]' in line:
                        if personal_info.get('hiringManagerName'):
                            line = f"Dear {personal_info['hiringManagerName']},"
                        else:
                            line = "Dear Hiring Manager,"
                    processed_lines.append(line)
            
            for line in processed_lines:
                para = doc.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if line.endswith(',') and (line.startswith('Dear ') or line == 'Sincerely,'):
                    doc.add_paragraph()
            
            byte_io = BytesIO()
            doc.save(byte_io)
            byte_io.seek(0)

            return send_file(
                byte_io,
                download_name="Cover_Letter.docx",
                as_attachment=True,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            logger.error(f"Error creating document: {str(e)}")
            return jsonify({"error": "Failed to create document"}), 500
    
    except Exception as e:
        logger.error(f"Error in generate_cover_letter: {e}")
        return jsonify({"error": "Internal server error"}), 500
    
@app.route("/", methods=["GET"])
def root():
    return jsonify({
        "message": "Cover Letter Generator API", 
        "status": "running", 
        "version": "1.0.0",
        "endpoints": ["/improve_resume", "/generate_cover_letter", "/health"]
    })

@app.route("/health", methods=["GET"])
def health_check():
    try:
        # Test Cohere connection
        co.chat(model=MODEL_NAME, message="test", max_tokens=1)
        return jsonify({"status": "healthy", "model": MODEL_NAME, "timestamp": datetime.now().isoformat()})
    except Exception as e:
        logger.error(f"Health check failed: {str(e)}")
        return jsonify({"status": "unhealthy", "error": "AI service unavailable"}), 503

if __name__ == "__main__":
    print("Starting Flask server...")
    print(f"Environment: {env}")
    print(f"Using Cohere key: {'✓ Set' if COHERE_API_KEY else '✗ Missing'}")
    
    # Don't run in debug mode in production
    debug_mode = env != "production"
    app.run(debug=debug_mode, host='0.0.0.0', port=int(os.getenv('PORT', 5000)))